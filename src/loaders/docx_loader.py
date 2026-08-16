from pathlib import Path
from uuid import UUID

from docx import Document as DocxDocument

from src.models import Document


def load_docx_document(
    file_path: Path,
    document_id: UUID
) -> Document:

    doc = DocxDocument(file_path)

    text = "\n".join(
        paragraph.text
        for paragraph in doc.paragraphs
    )

    return Document(
        document_id=document_id,
        filename=file_path.name,
        content=text
    )